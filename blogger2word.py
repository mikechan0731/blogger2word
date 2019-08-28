# -*-coding: utf-8-*-
# editor: MikeChan
# email: m7807031@gmail.com
# license: BSD-3

# OS import
import os
import random
from datetime import datetime, date
import time
import re

# Parser import
import requests
from urllib import request
from bs4 import BeautifulSoup
from bs4 import UnicodeDammit

# Office-docx lib import
import docx
from docx.shared import Cm, Pt  #add unit for word
from docx.enum.text import WD_ALIGN_PARAGRAPH # deal with alignment
from docx.shared import RGBColor

# ===== TODO ===== #
# 1.
# 2.
# 3.


# ===== Page Structure ===== #
# Blogger -> YearMonth -> Page
# Page -> [Title, Date, MainContent]
# MainContent -> 


# ===== Global Variables ===== #
base_blogger_url = "http://lulwechange.blogspot.com/"
root_blogger_dir = "blogger/"
start_year = 2015
YearMonth_list =[]
path_url_list = []


test_url_01 = "http://lulwechange.blogspot.com/2018/12/1.html" # with pic
test_url_02 = "http://lulwechange.blogspot.com/2018/10/blog-post_21.html" # with pics & font layout
test_url_03 = "http://lulwechange.blogspot.com/2016/06/"


# ===== Global Function ===== #
def timeIt(func):
    def wrapper(*args, **kw):
        t1 = time.time()
        res = func(*args, **kw)
        t2 = time.time()
        print("** timeIt Report: {} took {:.2f}s".format(func.__name__, t2-t1))
    return wrapper


def return_date_in_str():
    today = date.today()
    d1 = today.strftime("%Y%m%d")
    return d1


def save_html_to_txt(url):
    target = url
    res = requests.get(target)
    if res.status_code == requests.codes.ok:
        print("Target OK.")
        with open("{}_target.txt".format(return_date_in_str()), "w") as f:
            f.write(res.text)
    else:
        print("Bad Target.")


def is_article_exists(url):
    target = url
    res = requests.get(target)
    if res.status_code == requests.codes.ok:
        #print("Target OK.")
        if u"找不到文章" in res.text or u"網誌頁面不存在" in res.text:
            return False
        else:
            return True
    else:
        #print("Bad Target.")
        return False



# ====== Functions ===== #
def blogger_to_YearMonth_dir(blogger_url):
    year_now = int(datetime.now().year) 
    year_range = [year for year in range(start_year, year_now+1)]
    month_range = [month for month in range(1,12+1)]

    # create folder by YearMonth
    for year in year_range:
        for month in month_range:
            target_url = blogger_url+"{}/{:0>2d}/".format(year, month)
            if is_article_exists(target_url):
                print("{}{:0>2d}".format(year,month) + " OK! ")

                #create YearMonth folder
                if not os.path.exists(root_blogger_dir+"{}{:0>2d}".format(year,month)):
                    os.makedirs(root_blogger_dir+"{}{:0>2d}".format(year,month))

                YearMonth_list.append(target_url)

            else:
                print(target_url + " not exist")


def YearMonth_url_to_page_txt(YearMonth_url):
    page_list = []
    url_split = YearMonth_url.strip().split("/")
    year = url_split[-3]
    month = url_split[-2]
    
    target = YearMonth_url
    res = requests.get(target)
    if res.status_code == requests.codes.ok:
        with open("tmp/YM_tmp.txt", "w", encoding="utf-8") as f:
            f.write(res.text)
        
        with open("tmp/YM_tmp.txt", "r", encoding="utf-8") as r:
            count = 0
            for line in r:
                if u"閱讀更多" in line:
                    soup_line = BeautifulSoup(line, 'lxml')
                    a_tag = soup_line.find("a")
                    a_href = a_tag.get("href").split("#")[0]
                    a_title = a_tag.get("title")

                    #print(a_title, a_href)
                    count += 1
                    
                    page_title_url_str = a_title + "\t" + a_href

                    page_list.append(page_title_url_str)
                    
        
        with open("BloggerList.txt", "a", encoding="utf-8") as f:
            f.write("====="+ " " +year + month + " " +  u"共" + str(count) + u"篇" + "=====" +"\n")
            for page_info in page_list:
                f.write(page_info + "\n")
            f.write("\n")
    return


def BloggerList_parsing(BloggerList_path):
    with open(BloggerList_path, "r", encoding="utf-8") as r:
        save_path_now = ""
        for line in r:
            if line.startswith("====="):
                YM_path = line.split(" ")[1]
                save_path_now = root_blogger_dir + YM_path + "/"
            elif line.strip() == "":
                continue
            else:
                path_url_list.append([save_path_now, line.split("\t")[1].strip()])
    #print(path_url_list)



def single_page_to_content(save_path, page_url):
    target = page_url
    res = requests.get(target)
    if res.status_code == requests.codes.ok:
        soup = BeautifulSoup(res.text, "lxml")

        # page basic info
        # title
        blogger_title = soup.title.text.split(":")[0].strip()
        page_title = soup.title.text.split(":")[1].strip()
        print(page_title)

        # date
        date_tag = soup.findAll("h2", class_ = "date-header")
        date_str = date_tag[0].text
        date_list =[s for s in date_tag[0].text.split(" ")[0]]    

        date_arr = []
        s = ""
        for i in date_list:
            if i.isdigit():
                s += i
            else:
                date_arr.append(s)
                s = ""   
        date = "{}{:0>2d}{:0>2d}".format(int(date_arr[0]), int(date_arr[1]), int(date_arr[2]))
        print(date)
        
        # tag
        tags_tag = soup.findAll(rel="tag")
        tags_list = []
        for i in range(len(tags_tag)):
            tags_list.append(tags_tag[i].text)
        print(tags_list)

        # main content in page
        main_tag = soup.findAll(id="main")
        #main_text = main_tag[0].text


        # save to txt
        with open("tmp/tmp.txt", "w", encoding='utf-8') as f:
            f.write(str(main_tag[0]))

        # read txt, tranfer, and save to docx
        docx_name = date + "_" + page_title + ".docx"
        checked_docx_name = docx_name.replace("\\","-").replace("/","-").replace("|","-").replace("?","")\
            .replace('\"',"").replace("*","-").replace(":","-").replace("<","-").replace(">","-")

        txt_line_to_docx("tmp/tmp.txt", save_path, checked_docx_name, page_title, date_str, tags_list)

        # Done
        print("Page: " + date + "_" + page_title + " -> Saved.")
        print("= = =")
    else:
        print("Get " + page_url + " Failed!!")


# read txt, tranfer, and save to docx
def txt_line_to_docx(txt_name, docx_path, docx_name, title, date, tags_list):
    
    with open(txt_name, "r", encoding='utf-8') as r:
        # check line style and write to docx
        doc = docx.Document()
        doc.add_heading(title, level=1)
        doc.add_heading(date, level=3)
        tags_str = ""
        for tag in tags_list:
            tags_str += "#" + tag + "; "
        doc.add_heading(tags_str, level=3)
        doc.add_paragraph("")

        for line in r:
            # end of main content
            if u"張貼者" in line:
                break
            # get rid of repeated header info    
            if "date-header" in line or "entry-title" in line or "post-header" in line:
                continue
            if line.strip() == title.strip():
                continue

            soup_line = BeautifulSoup(line, "lxml")
            
            # check if href in line
            if soup_line.findAll('a', href=True):
                text = soup_line.text.strip()
                url = soup_line.findAll('a', href=True)[0]['href']

                # img
                if url.split(".")[-1].lower() == "jpg" or url.split(".")[-1].lower() == "png" or \
                    url.split(".")[-1].lower() == "bmp" or url.split(".")[-1].lower() == "jpeg":
                    request.urlretrieve(url,"tmp/tmp.jpg")
                    doc.add_picture("tmp/tmp.jpg", width=Cm(11))
                    last_paragraph = doc.paragraphs[-1] 
                    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER        

                # hyper-link
                else:
                    ok_line = text + "[" + url + "]"
                    doc.add_paragraph(ok_line)
            
            # normal text
            else:
                if soup_line.text.strip() == "":
                    continue 
                else:

                    if "<tr>" in line:
                        paragraph = doc.add_paragraph(soup_line.text)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        paragraph
                    else:
                         paragraph = doc.add_paragraph(soup_line.text)
        
        # save docx
        doc.save(docx_path + docx_name)




# ====== main() ===== #
def main():
    print("===== Blogger2Word Start! =====")
    
    if not os.path.exists('tmp'):
        os.makedirs('tmp')
    if not os.path.exists(root_blogger_dir):
        os.makedirs(root_blogger_dir)
    if os.path.exists("BloggerList.txt"):
        os.remove("BloggerList.txt")
    print("** Root Diretory Created.")

    #=== Test ===#
    #save_html_to_txt("http://lulwechange.blogspot.com/2015/08/")
    #is_article_exists("http://lulwechange.blogspot.com/2091/13/")

    #=== Main ===#

    # get YearMonth list
    print("** Year-Month List Checking...")
    blogger_to_YearMonth_dir(base_blogger_url) # return YM_url_list

    print("** BloggerList.txt Outputting...")
    for YM_url in YearMonth_list:
        YearMonth_url_to_page_txt(YM_url) #return txt 

    print("** BloggerList.txt Parsing...")
    BloggerList_parsing("BloggerList.txt") # return path_url_list
  

    print("** .docx Saving...")
    for path_url in path_url_list:
        single_page_to_content(path_url[0], path_url[1])



    print("===== Blogger2Word Done! ======")



if __name__=="__main__":
    main()